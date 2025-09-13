import streamlit as st
import json
import pandas as pd
import re
from io import BytesIO
from day6_grader import evaluate, companion_feedback
import fitz  # PyMuPDF for PDFs
import docx  # python-docx for DOCX
import os

# Google Drive API
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from googleapiclient.errors import HttpError
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from google.oauth2 import service_account

st.set_page_config(page_title="AI Grading App", layout="wide")

# ---------------------------
# Google Drive Helpers (Updated)
# ---------------------------
SERVICE_ACCOUNT_FILE = 'service_account.json'
SCOPES = ['https://www.googleapis.com/auth/drive']
model_id = "microsoft/phi-3.5-mini-instruct"

#https://www.googleapis.com/auth/drive https://www.googleapis.com/auth/drive.readonly

# def get_drive_service():
#     """Authenticate and return a Google Drive API service object."""
#     creds = None
#     token_file = 'token.json'

#     # Try to load existing credentials
#     if os.path.exists(token_file):
#         try:
#             creds = Credentials.from_authorized_user_file(token_file, SCOPES)
#         except Exception as e:
#             st.warning(f"⚠️ Ignoring corrupted token.json: {e}")
#             creds = None

#     # If no valid credentials, do manual auth
#     if not creds or not creds.valid:
#         if creds and creds.expired and creds.refresh_token:
#             creds.refresh(Request())
#         else:
#             if not os.path.exists("client_secrets.json"):
#                 st.error("⚠️ Please upload client_secrets.json to authenticate with Google Drive.")
#                 return None

#             flow = InstalledAppFlow.from_client_secrets_file('client_secrets.json', SCOPES)

#             # Get URL for manual authorization
#             auth_url, _ = flow.authorization_url(prompt='consent')
#             st.warning(f"🔑 Please go to this URL and authorize: {auth_url}")
#             auth_code = st.text_input("Paste the authorization code here:")

#             if auth_code:
#                 try:
#                     flow = InstalledAppFlow.from_client_secrets_file('client_secrets.json', SCOPES)

# # Works in Colab/Streamlit without a redirect URI issue
#                     creds = flow.run_console()

# # Save token
#                     with open(token_file, 'w') as token:
#                       token.write(creds.to_json())

#                     st.success("✅ Google Drive authenticated successfully!")


#                     st.success("✅ Google Drive authenticated successfully!")
#                 except Exception as e:
#                     st.error(f"❌ Authentication failed: {e}")
#                     return None
#             else:
#                 return None  # Stop if no auth code yet

#     try:
#         service = build('drive', 'v3', credentials=creds)
#         return service
#     except Exception as e:
#         st.error(f"⚠️ Could not authenticate Google Drive: {e}")
#         return None

def get_drive_service():
    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES
    )
    service = build('drive', 'v3', credentials=creds)
    return service

# Example: List first 10 files in the shared folder
service = get_drive_service()
results = service.files().list(
    pageSize=10, fields="files(id, name)"
).execute()
items = results.get('files', [])

if not items:
    print("No files found.")
else:
    for item in items:
        print(f"{item['name']} ({item['id']})")


def download_drive_file(service, file_id):
    """Download a file from Google Drive using its file ID."""
    try:
        # Get file metadata
        file = service.files().get(fileId=file_id, fields="id, name, mimeType").execute()
        file_name = file["name"]
        mime_type = file["mimeType"]

        fh = BytesIO()

        if mime_type == "application/vnd.google-apps.folder":
            # It's a folder, not a file
            return None, None, "Cannot download: Selected item is a folder."

        elif mime_type.startswith("application/vnd.google-apps."):
            # Google Docs/Sheets/Slides -> Export as PDF
            export_mime = "application/pdf"
            request = service.files().export_media(fileId=file_id, mimeType=export_mime)
            file_name += ".pdf"
        else:
            # Normal file download (PDF, DOCX, etc.)
            request = service.files().get_media(fileId=file_id)

        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
            if status:
                print(f"Downloading {file_name}: {int(status.progress() * 100)}%")

        fh.seek(0)
        return fh, file_name, None  # Success

    except Exception as e:
        return None, None, str(e)






def list_drive_files(service, page_size=20):
    """List files from Google Drive."""
    try:
        results = service.files().list(
            pageSize=page_size,
            fields="files(id, name)"
        ).execute()
        items = results.get('files', [])
        if not items:
            st.warning("⚠️ No files found in Google Drive.")
            return []
        return [(f['name'], f['id']) for f in items]
    except Exception as e:
        st.error(f"⚠️ Could not list files: {e}")
        return []

# ---------------------------
# File Parsing Helpers
# ---------------------------
def pdf_to_text(file):
    try:
        text = ""
        pdf_data = file.read()
        pdf_file = BytesIO(pdf_data)
        doc = fitz.open(stream=pdf_file, filetype="pdf")
        for page in doc:
            text += page.get_text() + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"❌ PDF extraction failed: {e}")
        return ""


def docx_to_text(file):
    try:
        doc_obj = docx.Document(file)
        return "\n".join([p.text for p in doc_obj.paragraphs if p.text.strip()]).strip()
    except Exception as e:
        st.error(f"❌ DOCX extraction failed: {e}")
        return ""


def detect_question(line):
    keywords = [
        r'\bdefine\b', r'\bdescribe\b', r'\bshow\b', r'\billustrate\b',
        r'\belaborate\b', r'\bexplain\b', r'\bgive\b',
        r'\bwho\b', r'\bwhat\b', r'\bwhere\b', r'\bwhen\b', r'\bwhy\b', r'\bhow\b'
    ]
    pattern = r'(\?$|:\s*$|' + "|".join(keywords) + r')'
    return re.search(pattern, line.strip(), flags=re.IGNORECASE)


def smart_parse_text_to_json(raw_text):
    raw_text = re.sub(r'\n+', '\n', raw_text.strip())
    lines = raw_text.split("\n")
    questions, current_q, current_a = [], None, []

    for line in lines:
        line = line.strip()
        if not line:
            continue
        if detect_question(line) or re.match(r'^\d+[\).]', line):
            if current_q:
                questions.append({
                    "question_id": f"Q{len(questions)+1}",
                    "question": current_q,
                    "student_answer": " ".join(current_a).strip(),
                    "correct_answer": "",
                    "max_score": 5
                })
            current_q, current_a = line, []
        else:
            current_a.append(line)
    if current_q:
        questions.append({
            "question_id": f"Q{len(questions)+1}",
            "question": current_q,
            "student_answer": " ".join(current_a).strip(),
            "correct_answer": "",
            "max_score": 5
        })
    return questions


# ---------------------------
# Sidebar Navigation
# ---------------------------
mode = st.sidebar.radio("Choose Mode", ["Grading Mode", "Companion Mode"])
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from io import BytesIO

# ---------------------------
# HELPER: Generate DOCX
# ---------------------------
def generate_docx(results, max_score):
    doc = Document()
    doc.add_heading("Graded Results", 0)

    for idx, q in enumerate(results, start=1):
        doc.add_paragraph(f"Q{idx}: {q['question']}")
        doc.add_paragraph(f"Student Answer: {q['student_answer']}")
        doc.add_paragraph(f"Correct Answer: {q['correct_answer']}")
        doc.add_paragraph(f"Model Score: {q['model_score']}")
        doc.add_paragraph(f"Final Score: {q['final_score']} / {max_score}")
        doc.add_paragraph(f"Feedback: {q.get('feedback', 'No feedback')}")
        doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------
# HELPER: Generate PDF
# ---------------------------
def generate_pdf(results, max_score):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Graded Results", styles["Title"]))
    story.append(Spacer(1, 20))

    for idx, q in enumerate(results, start=1):
        story.append(Paragraph(f"Q{idx}: {q['question']}", styles["Heading3"]))
        story.append(Paragraph(f"Student Answer: {q['student_answer']}", styles["Normal"]))
        story.append(Paragraph(f"Correct Answer: {q['correct_answer']}", styles["Normal"]))
        story.append(Paragraph(f"Model Score: {q['model_score']}", styles["Normal"]))
        story.append(Paragraph(f"Final Score: {q['final_score']} / {max_score}", styles["Normal"]))
        story.append(Paragraph(f"Feedback: {q.get('feedback', 'No feedback')}", styles["Normal"]))
        story.append(Spacer(1, 12))

    doc.build(story)
    buffer.seek(0)
    return buffer


# ---------------------------
# PAGE 1: GRADING MODE
# ---------------------------
if mode == "Grading Mode":
    st.title("📚 Xaminai-AI Grading Mode")
    difficulty = st.selectbox("Select grading difficulty:", ["Easy", "Medium", "Hard"])

    max_score = st.number_input("Set Max Score per Question", min_value=1, max_value=20, value=5, step=1)

    upload_option = st.radio("Choose Input Mode:",["✍️ Enter Text", "📂 Upload File", "☁️ Google Drive"],horizontal=True)
    correct_answers_text = None
    uploaded_file = None
    file_name = None
    correct_file = None


    if upload_option == "✍️ Enter Text":
      st.subheader("Enter Question and Answers")
      question = st.text_area("📝 Question")
      student_answer = st.text_area("👩‍🎓 Student Answer")
      correct_answer = st.text_area("✅ Correct Answer (optional)")

      if st.button("Grade This Answer"):
          if not question or not student_answer:
              st.warning("Please enter both question and student answer.")
          else:
              data = [{
                  "question_id": "Q1",
                  "question": question,
                  "student_answer": student_answer,
                  "correct_answer": correct_answer,
                  "grading_mode": difficulty
              }]

              # Save the input locally
              with open("uploaded.json", "w", encoding="utf-8") as f:
                  json.dump(data, f, indent=2, ensure_ascii=False)

              # Spinner while grading
              with st.spinner("Grading in progress..."):
                  evaluate("uploaded.json", "graded_results.json", difficulty=difficulty, max_score=max_score)

              # Read the results and display
              try:
                  with open("graded_results.json", "r", encoding="utf-8") as f:
                      results = json.load(f)

                  # Convert to DataFrame for summary
                  import pandas as pd
                  df = pd.DataFrame(results)

                  # Create 'feedback_short' if missing
                  if "feedback_short" not in df.columns:
                      df["feedback_short"] = df["feedback"].apply(lambda x: x[:100] + "..." if x else "")

                  # Grading Summary
                  summary_cols = ["question", "student_answer", "correct_answer", "model_score", "final_score", "feedback_short"]
                  st.subheader("🏷️ Grading Summary")
                  st.dataframe(df[summary_cols], use_container_width=True)

                  # Detailed Feedback
                  st.subheader("🔎 Detailed Feedback")
                  for idx, q in enumerate(results, start=1):
                      q_text = q.get("question", f"Question {idx}")
                      with st.expander(f"Question {idx}: {q_text}", expanded=False):
                          st.markdown(f"**Student Answer:** {q.get('student_answer', '')}")
                          st.markdown(f"**Correct Answer:** {q.get('correct_answer', '')}")
                          st.markdown(f"**Model Score:** {q.get('model_score', '')}  —  **Final Score:** {q.get('final_score', '')}")
                          st.markdown(f"**Feedback:** {q.get('feedback', 'No feedback available')}")
                          if q.get("improvement_steps"):
                              st.markdown("**🚀 Steps to Improve:**")
                              for step in q.get("improvement_steps", []):
                                  st.markdown(f"- {step}")
                          if q.get("keywords"):
                              st.markdown(f"**🔑 Keywords:** {', '.join(q.get('keywords', []))}")
                          if q.get("rule_score") is not None:
                              st.markdown(f"**Rule-based Score:** {q.get('rule_score')}")

              except Exception as e:
                  st.error(f"Could not load graded results: {e}")




    elif upload_option == "📂 Upload File":
      st.subheader("Upload Student Answers File")
      uploaded_file = st.file_uploader("Upload PDF/DOCX/JSON", type=["pdf", "docx", "json"])

      st.subheader("Optional: Upload Correct Answers File")
      correct_file = st.file_uploader("Upload Correct Answers (PDF/DOCX)", type=["pdf", "docx"])

      if uploaded_file:
          file_name = uploaded_file.name

          # Parse student answers
          if file_name.endswith(".pdf"):
              raw_text = pdf_to_text(uploaded_file)
              data = smart_parse_text_to_json(raw_text)
          elif file_name.endswith(".docx"):
              raw_text = docx_to_text(uploaded_file)
              data = smart_parse_text_to_json(raw_text)
          elif file_name.endswith(".json"):
              uploaded_file.seek(0)
              data = json.load(uploaded_file)
          else:
              st.error("Unsupported file format.")
              st.stop()

          # Parse optional correct answers file
          if correct_file:
              if correct_file.name.endswith(".pdf"):
                  correct_answers_text = pdf_to_text(correct_file)
              elif correct_file.name.endswith(".docx"):
                  correct_answers_text = docx_to_text(correct_file)

              correct_answers = [line.strip() for line in correct_answers_text.split("\n") if line.strip()]

              # attach correct answers to data
              for i, q in enumerate(data):
                  if i < len(correct_answers):
                      q["correct_answer"] = correct_answers[i]

          st.success(f"✅ Loaded {len(data)} questions")

          with open("uploaded.json", "w", encoding="utf-8") as f:
              json.dump(data, f, indent=2, ensure_ascii=False)

          evaluate("uploaded.json", "graded_results.json", difficulty=difficulty, max_score=max_score)

          with open("graded_results.json", "r", encoding="utf-8") as f:
              results = json.load(f)

          # --- Display results in a friendly table ---
          df = pd.DataFrame(results)
          if "feedback" not in df.columns:
              df["feedback"] = ""
          df["feedback_short"] = df["feedback"].astype(str).apply(
              lambda x: x if len(x) <= 120 else x[:117] + "..."
          )

          st.subheader("🏷️ Grading Summary")
          st.dataframe(
              df[["question", "student_answer", "correct_answer", "model_score", "final_score", "feedback_short"]],
              use_container_width=True
          )

          st.subheader("🔎 Detailed Feedback")
          for idx, q in enumerate(results, start=1):
              q_text = q.get("question", f"Question {idx}")
              with st.expander(f"Question {idx}: {q_text}", expanded=False):
                  st.markdown(f"**Student Answer:** {q.get('student_answer', '')}")
                  st.markdown(f"**Correct Answer:** {q.get('correct_answer', '')}")
                  st.markdown(f"**Model Score:** {q.get('model_score', '')}  —  **Final Score:** {q.get('final_score', '')}")
                  st.markdown(f"**Feedback:** {q.get('feedback', 'No feedback available')}")
                  if q.get("improvement_steps"):
                      st.markdown("**🚀 Steps to Improve:**")
                      for step in q.get("improvement_steps", []):
                          st.markdown(f"- {step}")
                  if q.get("keywords"):
                      st.markdown(f"**🔑 Keywords:** {', '.join(q.get('keywords', []))}")
                  if q.get("rule_score") is not None:
                      st.markdown(f"**Rule-based Score:** {q.get('rule_score')}")

    elif upload_option == "☁️ Google Drive":
      service = get_drive_service()
      if service:
          files = list_drive_files(service)
          if not files:
              st.info("📂 No files found in your Google Drive folder. Please upload a file and refresh.")
          else:
              file_dict = {f[0]: f[1] for f in files}
              selected_file = st.selectbox("Select a file from Google Drive", list(file_dict.keys()))

              st.subheader("Optional: Upload Correct Answers File (from Drive)")
              files_correct = list_drive_files(service)
              correct_file_id = None
              correct_selected = None
              if files_correct:
                  file_dict_correct = {f[0]: f[1] for f in files_correct}
                  correct_selected = st.selectbox(
                      "Select Correct Answers File (optional)",
                      ["None"] + list(file_dict_correct.keys())
                  )
                  if correct_selected != "None":
                      correct_file_id = file_dict_correct[correct_selected]

              if st.button("Download & Evaluate from Drive"):
                  file_id = file_dict[selected_file]
                  fh, downloaded_name, error = download_drive_file(service, file_id)

                  if fh:
                      # Save & rewind
                      with open(downloaded_name, "wb") as f:
                          f.write(fh.read())
                      fh.seek(0)
                      uploaded_file = fh
                      file_name = downloaded_name
                      st.success(f"✅ Downloaded {downloaded_name} successfully!")

                      # Parse student answers
                      if file_name.endswith(".pdf"):
                          raw_text = pdf_to_text(uploaded_file)
                          data = smart_parse_text_to_json(raw_text)
                      elif file_name.endswith(".docx"):
                          raw_text = docx_to_text(uploaded_file)
                          data = smart_parse_text_to_json(raw_text)
                      elif file_name.endswith(".json"):
                          uploaded_file.seek(0)
                          data = json.load(uploaded_file)
                      else:
                          st.error("Unsupported file format.")
                          st.stop()

                      # If optional correct answers file chosen
                      if correct_file_id:
                          fh_correct, correct_name, _ = download_drive_file(service, correct_file_id)
                          if fh_correct:
                              if correct_name.endswith(".pdf"):
                                  correct_answers_text = pdf_to_text(fh_correct)
                              elif correct_name.endswith(".docx"):
                                  correct_answers_text = docx_to_text(fh_correct)

                              correct_answers = [line.strip() for line in correct_answers_text.split("\n") if line.strip()]
                              for i, q in enumerate(data):
                                  if i < len(correct_answers):
                                      q["correct_answer"] = correct_answers[i]

                      # Save & evaluate using selected difficulty/max_score
                      # Save & evaluate using selected difficulty/max_score
                      with open("uploaded.json", "w", encoding="utf-8") as f:
                          json.dump(data, f, indent=2, ensure_ascii=False)

                      evaluate("uploaded.json", "graded_results.json",
                              difficulty=difficulty, max_score=max_score)

                      with open("graded_results.json", "r", encoding="utf-8") as f:
                          results = json.load(f)

                      # --- Display results in a friendly table ---
                      df = pd.DataFrame(results)
                      if "feedback" not in df.columns:
                          df["feedback"] = ""
                      df["feedback_short"] = df["feedback"].astype(str).apply(
                          lambda x: x if len(x) <= 120 else x[:117] + "..."
                      )

                      st.subheader("🏷️ Grading Summary")
                      st.dataframe(
                          df[["question", "student_answer", "correct_answer", "model_score", "final_score", "feedback_short"]],
                          use_container_width=True
                      )

                      st.subheader("🔎 Detailed Feedback")
                      for idx, q in enumerate(results, start=1):
                          q_text = q.get("question", f"Question {idx}")
                          with st.expander(f"Question {idx}: {q_text}", expanded=False):
                              st.markdown(f"**Student Answer:** {q.get('student_answer', '')}")
                              st.markdown(f"**Correct Answer:** {q.get('correct_answer', '')}")
                              st.markdown(f"**Model Score:** {q.get('model_score', '')}  —  **Final Score:** {q.get('final_score', '')}")
                              st.markdown(f"**Feedback:** {q.get('feedback', 'No feedback available')}")
                              if q.get("improvement_steps"):
                                  st.markdown("**🚀 Steps to Improve:**")
                                  for step in q.get("improvement_steps", []):
                                      st.markdown(f"- {step}")
                              if q.get("keywords"):
                                  st.markdown(f"**🔑 Keywords:** {', '.join(q.get('keywords', []))}")
                              if q.get("rule_score") is not None:
                                  st.markdown(f"**Rule-based Score:** {q.get('rule_score')}")




                          # Process file
                      if uploaded_file and file_name:
                          st.success(f"✅ Loaded: {file_name}")

                          # --- Reset pointer and parse student file ---
                          try:
                              uploaded_file.seek(0)
                          except Exception:
                              pass

                          raw_text = None
                          data = None
                          if file_name.endswith(".pdf"):
                              raw_text = pdf_to_text(uploaded_file)
                              data = smart_parse_text_to_json(raw_text)
                          elif file_name.endswith(".docx"):
                              raw_text = docx_to_text(uploaded_file)
                              data = smart_parse_text_to_json(raw_text)
                          elif file_name.endswith(".json"):
                              data = json.load(uploaded_file)
                          else:
                              st.error("❌ Unsupported file format.")
                              st.stop()

                          if not data:
                              st.error("⚠️ No questions found in the uploaded file.")
                              st.stop()

                          st.write(f"📄 Found **{len(data)} questions** in file.")
                          if st.checkbox("🔍 Show raw extracted text") and raw_text:
                              st.text_area("Extracted Text", raw_text, height=200)

                          # --- Parse optional correct answers file ---
                          correct_answers_json_path = None
                          if correct_file:
                              try:
                                  correct_file.seek(0)
                              except Exception:
                                  pass

                              if correct_file.name.endswith(".pdf"):
                                  correct_answers_text = pdf_to_text(correct_file)
                              elif correct_file.name.endswith(".docx"):
                                  correct_answers_text = docx_to_text(correct_file)
                              else:
                                  st.error("❌ Unsupported correct answers file format.")
                                  st.stop()

                              correct_answers_list = []
                              for i, line in enumerate(correct_answers_text.split("\n")):
                                  line = line.strip()
                                  if line:
                                      correct_answers_list.append({
                                          "question_id": f"Q{i+1}",
                                          "correct_answer": line
                                      })

                              correct_answers_json_path = "correct_answers.json"
                              with open(correct_answers_json_path, "w", encoding="utf-8") as f:
                                  json.dump(correct_answers_list, f, indent=2, ensure_ascii=False)

                              # Attach correct answers to student data in memory
                              for q in data:
                                  for ca in correct_answers_list:
                                      if q["question_id"] == ca["question_id"]:
                                          q["correct_answer"] = ca["correct_answer"]

                          # --- Save student answers for evaluation ---
                          input_path, output_path = "uploaded.json", "graded_results.json"
                          with open(input_path, "w", encoding="utf-8") as f:
                              json.dump(data, f, indent=2, ensure_ascii=False)

                          # --- Run evaluation ---
                          if correct_answers_json_path:
                              evaluate(input_path, output_path, difficulty=difficulty, max_score=max_score, correct_answers_file=correct_answers_json_path)
                          else:
                              evaluate(input_path, output_path, difficulty=difficulty, max_score=max_score)

                          # --- Load results and display ---
                          with open(output_path, "r", encoding="utf-8") as f:
                              results = json.load(f)

                          df = pd.DataFrame(results)
                          if "feedback" not in df.columns:
                              df["feedback"] = ""
                          df["feedback_short"] = df["feedback"].astype(str).apply(
                              lambda x: x if len(x) <= 120 else x[:117] + "..."
                          )

                          st.subheader("🏷️ Grading Summary")
                          st.dataframe(
                              df[["question", "student_answer", "correct_answer", "model_score", "final_score", "feedback_short"]],
                              use_container_width=True
                          )

                          st.subheader("🔎 Detailed Feedback")
                          for idx, q in enumerate(results, start=1):
                              q_text = q.get("question", f"Question {idx}")
                              with st.expander(f"Question {idx}: {q_text}", expanded=False):
                                  st.markdown(f"**Student Answer:** {q.get('student_answer', '')}")
                                  st.markdown(f"**Correct Answer:** {q.get('correct_answer', '')}")
                                  st.markdown(f"**Model Score:** {q.get('model_score', '')}  —  **Final Score:** {q.get('final_score', '')}")
                                  st.markdown(f"**Feedback:** {q.get('feedback', 'No feedback available')}")
                                  if q.get("improvement_steps"):
                                      st.markdown("**🚀 Steps to Improve:**")
                                      for step in q.get("improvement_steps", []):
                                          st.markdown(f"- {step}")
                                  if q.get("keywords"):
                                      st.markdown(f"**🔑 Keywords:** {', '.join(q.get('keywords', []))}")
                                  if q.get("rule_score") is not None:
                                      st.markdown(f"**Rule-based Score:** {q.get('rule_score')}")

                          st.subheader("📥 Export Results")
                          try:
                              docx_buffer = generate_docx(results, max_score)
                              st.download_button(
                                  label="📘 Download DOCX Report",
                                  data=docx_buffer,
                                  file_name="graded_results.docx",
                                  mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                              )
                          except Exception as e:
                              st.warning(f"⚠️ DOCX generation skipped: {e}")

                          try:
                              pdf_buffer = generate_pdf(results, max_score)
                              st.download_button(
                                  label="📄 Download PDF Report",
                                  data=pdf_buffer,
                                  file_name="graded_results.pdf",
                                  mime="application/pdf"
                              )
                          except Exception as e:
                              st.warning(f"⚠️ PDF generation skipped: {e}")




# ---------------------------
# PAGE 2: COMPANION MODE
# ---------------------------
elif mode == "Companion Mode":
    st.title("🤝 AI Companion Tutor")
    st.write("Get guided feedback on your answers and learn how to improve!")

    upload_option = st.radio("Choose Input Method", ["✏️ Manual Input", "📂 Upload File"])

    question = ""
    student_answer = ""
    correct_answer = ""

    if upload_option == "✏️ Manual Input":
        question = st.text_area("Enter your Question")
        student_answer = st.text_area("Enter your Answer")
        correct_answer = st.text_area("Correct Answer (optional, for better guidance)")

    elif upload_option == "📂 Upload File":
        file = st.file_uploader("Upload a JSON, PDF, or DOCX file", type=["json", "pdf", "docx"])
        if file:
            st.success(f"✅ Uploaded: {file.name}")
            if file.name.endswith(".pdf"):
                raw_text = pdf_to_text(file)
                parsed = smart_parse_text_to_json(raw_text)
            elif file.name.endswith(".docx"):
                raw_text = docx_to_text(file)
                parsed = smart_parse_text_to_json(raw_text)
            elif file.name.endswith(".json"):
                parsed = json.load(file)
            else:
                parsed = []

            if parsed:
                q_idx = st.number_input("Select Question Index", min_value=1, max_value=len(parsed), value=1)
                selected = parsed[q_idx - 1]
                question = selected.get("question", "")
                student_answer = selected.get("student_answer", "")
                correct_answer = selected.get("correct_answer", "")

                st.write(f"**Question:** {question}")
                st.write(f"**Student Answer:** {student_answer}")

    max_score = st.number_input("Max Score", min_value=1, max_value=10, value=5)

    if st.button("Get Guidance"):
        if not question or not student_answer:
            st.warning("Please provide a question and your answer.")
        else:
            with st.spinner("Generating feedback..."):
                result = companion_feedback(question, student_answer, correct_answer, max_score)

            feedback = result.get("feedback", "").replace(question, "").replace(student_answer, "")
            st.subheader("📢 Feedback")
            st.write(feedback if feedback else "No feedback available")

            keywords = result.get("keywords", [])
            st.subheader("🔑 Keywords for a Perfect Answer")
            st.write(", ".join(keywords) if keywords else "No keywords found")

            steps = result.get("improvement_steps", [])
            st.subheader("🚀 Steps to Improve")
            if steps:
                for step in steps:
                    st.markdown(f"- {step}")
            else:
                st.write("No improvement steps available")



